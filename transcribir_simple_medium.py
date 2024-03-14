import os
import whisper
from docx import Document

# Get filenames in the current directory
filenames = os.listdir()

model = whisper.load_model("medium")

text = ''

# Print each filename
for filename in filenames:
    try:
        result = model.transcribe(filename)
        text += filename + '\n\n' + result["text"] + '\n\n\n'
    except:
        pass 


# Create a new document object
document = Document()

# Add a paragraph to the document
paragraph = document.add_paragraph()

# Add the string to the paragraph
paragraph.add_run(text)

# Save the document with a filename
document.save("my_document.docx")

